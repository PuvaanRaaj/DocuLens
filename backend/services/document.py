"""
Document generation service — produces DOCX, PDF, and TXT output
from a structured JSON document representation.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from typing import Dict, Any, List
import io


class DocumentService:
    """Generates output documents in multiple formats from structured JSON."""

    # ── DOCX ──────────────────────────────────────────────────────────

    def generate_docx(self, structure: Dict[str, Any]) -> io.BytesIO:
        """Generates a DOCX file from the structured JSON data."""
        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        elements = structure.get("elements", [])

        for element in elements:
            el_type = element.get("type", "paragraph")
            text = element.get("text", "")

            if el_type == "heading1":
                doc.add_heading(text, level=1)
            elif el_type == "heading2":
                doc.add_heading(text, level=2)
            elif el_type == "heading3":
                doc.add_heading(text, level=3)
            elif el_type == "bullet_list":
                items = element.get("items", [])
                for item in items:
                    doc.add_paragraph(item, style='List Bullet')
            elif el_type == "numbered_list":
                items = element.get("items", [])
                for item in items:
                    doc.add_paragraph(item, style='List Number')
            else:
                doc.add_paragraph(text)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # ── PDF ────────────────────────────────────────────────────────────

    def generate_pdf(self, structure: Dict[str, Any]) -> io.BytesIO:
        """Generates a PDF file from the structured JSON data using fpdf2."""
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Use built-in fonts (no external font files needed)
        pdf.set_font("Helvetica", size=11)

        elements = structure.get("elements", [])

        for element in elements:
            el_type = element.get("type", "paragraph")
            text = element.get("text", "")

            if el_type in ("heading1", "heading2", "heading3"):
                sizes = {"heading1": 22, "heading2": 18, "heading3": 14}
                pdf.set_font("Helvetica", style="B", size=sizes[el_type])
                pdf.ln(4)
                pdf.multi_cell(0, 10, text)
                pdf.ln(2)
                pdf.set_font("Helvetica", size=11)

            elif el_type == "bullet_list":
                items = element.get("items", [])
                for item in items:
                    pdf.set_x(pdf.l_margin + 10)
                    pdf.multi_cell(0, 7, f"\u2022  {item}")

            elif el_type == "numbered_list":
                items = element.get("items", [])
                for idx, item in enumerate(items, 1):
                    pdf.set_x(pdf.l_margin + 10)
                    pdf.multi_cell(0, 7, f"{idx}.  {item}")

            else:
                pdf.multi_cell(0, 7, text)
                pdf.ln(2)

        buffer = io.BytesIO()
        pdf.output(buffer)
        buffer.seek(0)
        return buffer

    # ── TXT ────────────────────────────────────────────────────────────

    def generate_txt(self, structure: Dict[str, Any]) -> io.BytesIO:
        """Generates a plain-text file from the structured JSON data."""
        lines: List[str] = []
        elements = structure.get("elements", [])

        for element in elements:
            el_type = element.get("type", "paragraph")
            text = element.get("text", "")

            if el_type == "heading1":
                lines.append(f"# {text}")
                lines.append("")
            elif el_type == "heading2":
                lines.append(f"## {text}")
                lines.append("")
            elif el_type == "heading3":
                lines.append(f"### {text}")
                lines.append("")
            elif el_type == "bullet_list":
                items = element.get("items", [])
                for item in items:
                    lines.append(f"  - {item}")
                lines.append("")
            elif el_type == "numbered_list":
                items = element.get("items", [])
                for idx, item in enumerate(items, 1):
                    lines.append(f"  {idx}. {item}")
                lines.append("")
            else:
                lines.append(text)
                lines.append("")

        content = "\n".join(lines)
        buffer = io.BytesIO()
        buffer.write(content.encode("utf-8"))
        buffer.seek(0)
        return buffer
